import torch
from transformers import AutoTokenizer, AutoModelForSequenceClassification
from sentence_transformers import SentenceTransformer, util
import pandas as pd
from sklearn.preprocessing import LabelEncoder
import os
import fnmatch
from docx import Document
import re

# Инициализация модели SentenceTransformer глобально
sentence_model = SentenceTransformer('all-MiniLM-L6-v2')

# 1. Загрузка модели и токенизатора
model = AutoModelForSequenceClassification.from_pretrained('saved_model')
tokenizer = AutoTokenizer.from_pretrained('saved_model')

df = pd.read_excel(r'C:\Users\selvl\Desktop\ITOG\Hackaton\dataset.xlsx')
df['Solution'] = df['Solution'].fillna('')

# 2. Подготовка LabelEncoder и классов
classes_path = 'saved_model/classes.txt'
if os.path.exists(classes_path):
    with open(classes_path, 'r', encoding='utf-8') as f:
        classes = [line.strip() for line in f.readlines()]
else:
    # Объединяем редкие классы так же, как при обучении
    threshold = 5
    class_counts = df['Label'].value_counts()
    rare_classes = class_counts[class_counts < threshold].index.tolist()
    df['label_combined'] = df['Label'].apply(
        lambda x: 'Other' if x in rare_classes else x
    )

    label_encoder = LabelEncoder()
    df['label_encoded'] = label_encoder.fit_transform(df['label_combined'])
    classes = label_encoder.classes_

    # Сохраняем классы в файл для будущего использования
    with open(classes_path, 'w', encoding='utf-8') as f:
        for cls in classes:
            f.write(f"{cls}\n")

# 3. Функция для классификации сервиса
def classify_service(text):
    inputs = tokenizer(
        text,
        return_tensors='pt',
        truncation=True,
        padding='max_length',
        max_length=128
    )
    outputs = model(**inputs)
    probabilities = torch.nn.functional.softmax(outputs.logits, dim=1)
    top_prob, top_class = torch.topk(probabilities, k=3)
    services = []
    for prob, cls in zip(top_prob[0], top_class[0]):
        cls = cls.item()
        services.append({
            'service': classes[cls],
            'probability': prob.item()
        })
    return services

all_topics = df['Topic'].tolist()
topic_embeddings = None

# Загружаем эмбеддинги обращений из файла, если они сохранены
embeddings_path = 'topic_embeddings.pt'
if os.path.exists(embeddings_path):
    topic_embeddings = torch.load(embeddings_path)
else:
    # Вычисляем эмбеддинги и сохраняем их
    topic_embeddings = sentence_model.encode(all_topics, convert_to_tensor=True)
    torch.save(topic_embeddings, embeddings_path)

# 5. Функция для поиска похожих обращений
def find_similar_topics(query, top_k=5):
    if topic_embeddings is None:
        print("Эмбеддинги обращений не загружены.")
        return []
    query_embedding = sentence_model.encode(query, convert_to_tensor=True)
    cos_scores = util.pytorch_cos_sim(query_embedding, topic_embeddings)[0]
    top_results = torch.topk(cos_scores, k=top_k)

    similar_topics = []
    for score, idx in zip(top_results.values, top_results.indices):
        idx = idx.item()
        similar_topics.append({
            'topic': df.iloc[idx]['Topic'],
            'solution': df.iloc[idx]['Solution'],
            'similarity': score.item()
        })
    return similar_topics

# 6. Подготовка данных для поиска релевантных инструкций
instructions_folder = r'C:\Users\selvl\Desktop\ITOG\Hackaton'

def clean_text(text):
    # Удаляем непечатаемые символы
    text = re.sub(r'[^\x20-\x7E]+', ' ', text)
    # Заменяем множественные пробелы одним пробелом
    text = re.sub(r'\s+', ' ', text).strip()
    return text

def load_instructions(folder_path):
    instructions = []
    try:
        files_in_directory = os.listdir(folder_path)
    except Exception as e:
        return instructions

    file_paths = []
    for filename in files_in_directory:
        if fnmatch.fnmatch(filename.lower(), '*.docx'):
            full_path = os.path.join(folder_path, filename)
            file_paths.append(full_path)
    if not file_paths:
        print("Папка с инструкциями пуста или файлы не найдены.")
        return instructions

    for file_path in file_paths:
        try:
            doc = Document(file_path)
            full_text = []

            # Извлечение текста из параграфов
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    full_text.append(text)

            # Извлечение текста из таблиц
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            text = paragraph.text.strip()
                            if text:
                                full_text.append(text)

            # Извлечение текста из заголовков и сносок
            for section in doc.sections:
                header = section.header
                for paragraph in header.paragraphs:
                    text = paragraph.text.strip()
                    if text:
                        full_text.append(text)

                footer = section.footer
                for paragraph in footer.paragraphs:
                    text = paragraph.text.strip()
                    if text:
                        full_text.append(text)

            # Объединяем текст и удаляем лишние переносы строк
            content = '\n'.join(full_text)
            content = re.sub(r'\n+', '\n', content)  # Удаляем повторяющиеся символы новой строки

            # Очищаем текст от непечатаемых символов и лишних пробелов
            content = clean_text(content)

            if content.strip():
                # Удаляем вывод пути к файлу
                # print(f"Извлечённый текст из {file_path}:\n{content[:200]}...\n")
                instructions.append({
                    'title': os.path.basename(file_path),
                    'content': content
                })
            else:
                print(f"Файл {file_path} содержит пустой текст и будет пропущен.")
        except Exception as e:
            print(f"Ошибка при чтении файла {file_path}: {e}")
    return instructions

instructions = load_instructions(instructions_folder)

instruction_texts = [instr['content'] for instr in instructions if instr['content'].strip()]

# Ограничение длины текстов
max_length = 512
instruction_texts = [text[:max_length] for text in instruction_texts]

instructions_embeddings_path = 'instruction_embeddings.pt'

# Удаление старого файла эмбеддингов
if os.path.exists(instructions_embeddings_path):
    os.remove(instructions_embeddings_path)

# Проверяем, есть ли тексты инструкций
if instruction_texts:
    # Вычисляем эмбеддинги и сохраняем их
    instruction_embeddings = sentence_model.encode(instruction_texts, convert_to_tensor=True)
    torch.save(instruction_embeddings, instructions_embeddings_path)
else:
    instruction_embeddings = None

if instruction_embeddings is not None:
    expected_shape = (len(instruction_texts), 384)
    if instruction_embeddings.size() == expected_shape:
        print(f"Размер эмбеддингов инструкций корректен: {instruction_embeddings.size()}")
    else:
        print(f"Некорректный размер эмбеддингов: {instruction_embeddings.size()}, ожидаемый: {expected_shape}")
else:
    print("Не удалось вычислить эмбеддинги инструкций.")

# 7. Функция для поиска релевантных инструкций
def find_relevant_instructions(query, top_k=3):
    print("Вызвана функция find_relevant_instructions")
    if instruction_embeddings is None or not instruction_embeddings.size(0):
        print("Инструкции отсутствуют или не загружены.")
        return []
    try:
        query_embedding = sentence_model.encode(query, convert_to_tensor=True)
        cos_scores = util.pytorch_cos_sim(query_embedding, instruction_embeddings)[0]
        top_k = min(top_k, instruction_embeddings.size(0))
        top_results = torch.topk(cos_scores, k=top_k)

        relevant_instructions = []
        for score, idx in zip(top_results.values, top_results.indices):
            idx = idx.item()
            relevant_instructions.append({
                'title': instructions[idx]['title'],
                'content': instructions[idx]['content'],
                'similarity': score.item()
            })
        return relevant_instructions
    except Exception as e:
        print(f"Ошибка в find_relevant_instructions: {e}")
        return []

responseTest = "Все работает"

# 8. Функция помощника
def assistant_response(query):
    # Определение сервиса
    services = classify_service(query)
    services_output = []
    for s in services:
        services_output.append({
            'service': s['service'],
            'probability': f"{s['probability']:.4f}"
        })

    # Поиск похожих обращений
    similar_topics = find_similar_topics(query)
    topics_output = []
    if similar_topics:
        for topic in similar_topics:
            solution = topic['solution'] if topic['solution'] else "отсутствует"
            topics_output.append({
                'similarity': f"{topic['similarity']:.4f}",
                'topic': topic['topic'],
                'solution': solution
            })
    else:
        topics_output = None

    # Поиск релевантных инструкций
    relevant_instructions = find_relevant_instructions(query)
    instructions_output = []
    if relevant_instructions:
        for instr in relevant_instructions:
            instructions_output.append({
                'similarity': f"{instr['similarity']:.4f}",
                'title': instr['title'],
                'content': instr['content']
            })
    else:
        instructions_output = None

    # Собираем все данные в словарь
    response_data = {
        'services': services_output,
        'similar_topics': topics_output,
        'relevant_instructions': instructions_output
    }

    return response_data

# 9. Запуск помощника
if __name__ == "__main__":
    query = input("Опишите вашу проблему: ")
    assistant_response(query)